from __future__ import annotations

import sys
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.shared import Inches, Pt


def ensure_styles(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(10.5)

    for style_name, size in [("Heading 1", 15), ("Heading 2", 13), ("Heading 3", 11)]:
        style = doc.styles[style_name]
        style.font.name = "Arial"
        style.font.size = Pt(size)

    if "Body Bullet" not in [s.name for s in doc.styles]:
        style = doc.styles.add_style("Body Bullet", WD_STYLE_TYPE.PARAGRAPH)
        style.base_style = doc.styles["Normal"]
        style.font.name = "Arial"
        style.font.size = Pt(10.5)


def set_page_layout(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)


def add_table(doc: Document, rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1, cols=len(rows[0]))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    hdr_cells = table.rows[0].cells
    for i, cell_text in enumerate(rows[0]):
        hdr_cells[i].text = cell_text.strip()

    for row in rows[1:]:
        cells = table.add_row().cells
        for i, cell_text in enumerate(row):
            cells[i].text = cell_text.strip()

    doc.add_paragraph("")


def parse_table_block(lines: list[str], start: int) -> tuple[list[list[str]], int]:
    rows: list[list[str]] = []
    idx = start
    while idx < len(lines) and lines[idx].strip().startswith("|"):
        raw = lines[idx].strip().strip("|")
        row = [part.strip() for part in raw.split("|")]
        rows.append(row)
        idx += 1

    if len(rows) >= 2 and all(set(cell) <= {"-", ":"} for cell in rows[1]):
        rows.pop(1)

    return rows, idx


def render_markdown(src: Path, dst: Path) -> None:
    doc = Document()
    ensure_styles(doc)
    set_page_layout(doc)

    lines = src.read_text(encoding="utf-8").splitlines()
    idx = 0

    while idx < len(lines):
        line = lines[idx].rstrip()
        stripped = line.strip()

        if not stripped:
            idx += 1
            continue

        if stripped.startswith("# "):
            doc.add_heading(stripped[2:].strip(), level=0)
            idx += 1
            continue

        if stripped.startswith("## "):
            doc.add_heading(stripped[3:].strip(), level=1)
            idx += 1
            continue

        if stripped.startswith("### "):
            doc.add_heading(stripped[4:].strip(), level=2)
            idx += 1
            continue

        if stripped.startswith("|"):
            rows, idx = parse_table_block(lines, idx)
            add_table(doc, rows)
            continue

        if stripped.startswith("- "):
            doc.add_paragraph(stripped[2:].strip(), style="List Bullet")
            idx += 1
            continue

        doc.add_paragraph(stripped)
        idx += 1

    dst.parent.mkdir(parents=True, exist_ok=True)
    doc.save(dst)


def main() -> int:
    if len(sys.argv) != 3:
        print("Usage: python scripts/markdown_to_docx.py <input.md> <output.docx>")
        return 1

    src = Path(sys.argv[1]).resolve()
    dst = Path(sys.argv[2]).resolve()
    render_markdown(src, dst)
    print(dst)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
